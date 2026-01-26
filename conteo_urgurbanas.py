import streamlit as st
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.ticker as mtick
import matplotlib.dates as mdates
import plotly.express as px
import plotly.graph_objects as go
import tempfile
import os
from datetime import datetime
from io import BytesIO
import unicodedata
import base64
import re

# Configuración de la página
st.set_page_config(
    page_title="Analizador de Urgencias Operativas",
    page_icon="📊",
    layout="wide"
)

# --- FUNCIONES DE UTILIDAD ---

def formatear_periodo_es(periodo):
    if pd.isna(periodo): return ""
    meses = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril", 5: "Mayo", 6: "Junio", 
             7: "Julio", 8: "Agosto", 9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}
    try:
        return f"{meses.get(periodo.month, '')} {periodo.year}"
    except:
        return str(periodo)

def limpiar_texto(texto):
    # CORRECCIÓN: Si es nulo, devolvemos None para que pandas lo ignore en conteos
    if pd.isna(texto): return None
    if not isinstance(texto, str): return str(texto).lower().strip()
    return unicodedata.normalize('NFD', texto).encode('ascii', 'ignore').decode('utf-8').lower().strip()

def normalizar_genero(texto):
    t = limpiar_texto(texto)
    if not t: return None
    if t in ['hombre', 'masculino', 'm', 'varon']: return 'Masculino'
    if t in ['mujer', 'femenino', 'f', 'dama']: return 'Femenino'
    return None

def limpiar_y_categorizar_edad(valor):
    if pd.isna(valor): return None
    s = str(valor)
    numeros = re.findall(r'\d+', s)
    if not numeros: return None
    try: edad = int(numeros[0])
    except: return None
    if edad < 18: return "Menor (0-17)"
    if 18 <= edad <= 29: return "Joven (18-29)"
    if 30 <= edad <= 59: return "Adulto (30-59)"
    if edad >= 60: return "Mayor (60+)"
    return None

def parsear_fecha(fecha):
    if pd.isna(fecha): return None
    if isinstance(fecha, (datetime, pd.Timestamp)): return fecha
    fecha_str = str(fecha).strip()
    formatos = ['%d/%m/%Y', '%d-%m-%Y', '%Y-%m-%d', '%Y/%m/%d', '%d.%m.%Y', '%Y.%m.%d']
    for fmt in formatos:
        try: return datetime.strptime(fecha_str.split(' ')[0], fmt)
        except: continue
    try: return pd.to_datetime(fecha, dayfirst=True).to_pydatetime()
    except: return None

# --- FUNCIÓN: CLASIFICACIÓN DE ENFERMEDADES POR TEXTO ---
def clasificar_enfermedad(texto):
    t = limpiar_texto(texto)
    if not t: return "No especificado"
    
    keywords = {
        'Diabetes / Glucosa': ['diabet', 'glucos', 'azucar', 'hiperglucemia', 'hipoglucemia', 'insulin'],
        'Hipertensión / Presión': ['hiperten', 'presion', 't/a', 'hta', 'tension'],
        'Traumatismo / Caída': ['caida', 'golpe', 'trauma', 'herida', 'escalera', 'tropez', 'altura'],
        'Respiratorio': ['respiratori', 'disnea', 'aire', 'oxigeno', 'asm', 'epoc', 'bronqu'],
        'Neurológico (Convulsión/EVC)': ['convulsi', 'epilep', 'sincop', 'desmay', 'inconsciente', 'evc', 'cerebr'],
        'Cardíaco': ['infarto', 'cardiac', 'pecho', 'corazon', 'taquicardia'],
        'Gastrointestinal': ['dolor abdominal', 'estomac', 'vomito', 'diarrea', 'gastrit'],
        'Embarazo / Parto': ['embaraz', 'parto', 'labor', 'gestan', 'bebe'],
        'Intoxicación': ['intoxica', 'veneno', 'sustancia', 'alcohol', 'ebri'],
        'Violencia / Agresión': ['agresion', 'riña', 'golpead', 'arma']
    }
    
    found = []
    for cat, terms in keywords.items():
        for term in terms:
            if term in t:
                found.append(cat)
                break 
    
    if not found: return "Otros / No detectado"
    return ", ".join(found)

# --- FUNCIONES DE GRÁFICAS ---

def generar_grafica_bar(conteo, titulo, filename):
    if conteo is None or conteo.empty: return None
    df_plot = conteo.reset_index()
    df_plot.columns = ['Categoría', 'Cantidad']
    # Ignoramos nulos que se hayan colado como etiquetas
    df_plot = df_plot[df_plot['Categoría'].notna() & (df_plot['Categoría'] != 'None')]
    
    plt.figure(figsize=(10, 6))
    colors = plt.cm.viridis(np.linspace(0, 1, len(df_plot)))
    bars = plt.bar(df_plot['Categoría'].astype(str), df_plot['Cantidad'], color=colors)
    plt.title(titulo, fontsize=12, fontweight='bold')
    plt.xticks(rotation=45, ha='right', fontsize=8)
    for bar in bars:
        plt.text(bar.get_x() + bar.get_width()/2., bar.get_height() + 0.1,
                f'{int(bar.get_height())}', ha='center', va='bottom', fontsize=8)
    plt.tight_layout()
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

# (Se mantienen funciones generar_grafica_linea_... y generar_grafica_plotly_... del original)
# ... [Omitido por brevedad, se asume igual al original del usuario] ...

def generar_grafica_linea_simple(datos, titulo, xlabel, ylabel, filename):
    if datos.empty: return None
    df_plot = datos.reset_index()
    df_plot.columns = ['Periodo', 'Cantidad']
    try:
        df_plot['Fecha_X'] = df_plot['Periodo'].dt.to_timestamp()
    except:
        df_plot['Fecha_X'] = pd.to_datetime(df_plot['Periodo'].astype(str), errors='coerce')
    df_plot = df_plot.dropna(subset=['Fecha_X']).sort_values('Fecha_X')
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(df_plot['Fecha_X'], df_plot['Cantidad'], marker='o', linestyle='-', color='teal', linewidth=2)
    ax.set_title(titulo, fontsize=12, fontweight='bold')
    ax.set_xlabel(xlabel, fontweight='bold')
    ax.set_ylabel(ylabel, fontweight='bold')
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%B %Y'))
    ax.xaxis.set_major_locator(mdates.MonthLocator())
    ax.tick_params(axis='x', rotation=45, labelsize=8)
    ax.grid(True, alpha=0.3)
    plt.tight_layout()
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

def generar_grafica_linea_multiple(df_long, col_x, col_y, col_grupo, titulo, filename):
    if df_long.empty: return None
    fig, ax = plt.subplots(figsize=(12, 7))
    grupos = df_long[col_grupo].unique()
    cmap = plt.get_cmap('tab10')
    for i, grupo in enumerate(grupos):
        subset = df_long[df_long[col_grupo] == grupo].copy()
        try: subset['Fecha_X'] = subset[col_x].dt.to_timestamp()
        except: subset['Fecha_X'] = pd.to_datetime(subset[col_x].astype(str), errors='coerce')
        subset = subset.dropna(subset=['Fecha_X']).sort_values('Fecha_X')
        ax.plot(subset['Fecha_X'], subset[col_y], marker='o', linestyle='-', linewidth=2, label=grupo, color=cmap(i % 10))
    ax.set_title(titulo, fontsize=14, fontweight='bold')
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%B %Y'))
    ax.xaxis.set_major_locator(mdates.MonthLocator())
    ax.tick_params(axis='x', rotation=45, labelsize=8)
    ax.grid(True, alpha=0.3)
    ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left', borderaxespad=0., fontsize='small')
    plt.tight_layout()
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

def generar_grafica_linea_porcentaje_genero(df_long, col_periodo, col_pct, col_genero, titulo, filename):
    if df_long.empty: return None
    fig, ax = plt.subplots(figsize=(12, 7))
    color_map = {'Masculino': 'blue', 'Femenino': 'purple'}
    for genero in ['Masculino', 'Femenino']:
        subset = df_long[df_long[col_genero] == genero].copy()
        if subset.empty: continue
        try: subset['Fecha_X'] = subset[col_periodo].dt.to_timestamp()
        except: subset['Fecha_X'] = pd.to_datetime(subset[col_periodo].astype(str), errors='coerce')
        subset = subset.dropna(subset=['Fecha_X']).sort_values('Fecha_X')
        color = color_map.get(genero, 'grey')
        ax.plot(subset['Fecha_X'], subset[col_pct], marker='o', linestyle='-', linewidth=3, label=genero, color=color)
    ax.set_title(titulo, fontsize=14, fontweight='bold')
    ax.set_ylim(0, 115)
    ax.yaxis.set_major_formatter(mtick.PercentFormatter(decimals=0))
    ax.xaxis.set_major_formatter(mdates.DateFormatter('%B %Y'))
    ax.xaxis.set_major_locator(mdates.MonthLocator())
    ax.tick_params(axis='x', rotation=45, labelsize=9)
    ax.grid(True, alpha=0.3, axis='y')
    ax.legend(loc='best', fontsize='medium')
    plt.tight_layout()
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

def generar_grafica_circulos_edad_word(df_data, titulo, filename):
    if df_data.empty: return None
    fig, ax = plt.subplots(figsize=(12, 8))
    colonias = df_data['Colonia'].unique()
    rangos = ["Menor (0-17)", "Joven (18-29)", "Adulto (30-59)", "Mayor (60+)"]
    ax.set_xlim(-0.5, len(colonias) - 0.5)
    ax.set_ylim(-0.5, len(rangos) - 0.5)
    col_map = {c: i for i, c in enumerate(colonias)}
    rango_map = {r: i for i, r in enumerate(rangos)}
    for _, row in df_data.iterrows():
        c_idx = col_map.get(row['Colonia'])
        r_idx = rango_map.get(row['Rango'])
        if c_idx is not None and r_idx is not None:
            pct = row['Porcentaje']
            ax.scatter(c_idx, r_idx, s=pct * 30, c='black', alpha=0.9, zorder=3)
            ax.text(c_idx, r_idx, f"{pct:.0f}%", ha='center', va='center', fontsize=9, fontweight='bold', color='white', zorder=4)
    ax.set_xticks(range(len(colonias)))
    ax.set_xticklabels(colonias, rotation=45, ha='right')
    ax.set_yticks(range(len(rangos)))
    ax.set_yticklabels(rangos)
    ax.set_title(titulo, fontsize=14, fontweight='bold')
    ax.grid(True, alpha=0.2, linestyle='--', zorder=0)
    plt.tight_layout()
    path = os.path.join(tempfile.gettempdir(), filename)
    plt.savefig(path, dpi=200, bbox_inches='tight')
    plt.close()
    return path

def generar_grafica_plotly_bar(conteo, titulo):
    if conteo is None or conteo.empty: return px.bar(title="Sin datos")
    df_plot = conteo.reset_index()
    df_plot.columns = ['Categoría', 'Cantidad']
    df_plot = df_plot[df_plot['Categoría'].notna()]
    fig = px.bar(df_plot, x='Categoría', y='Cantidad', title=titulo, color='Cantidad')
    return fig

def generar_grafica_plotly_linea(df_long, col_periodo, col_y, col_color, titulo, es_porcentaje=False):
    if df_long.empty: return px.line(title="Sin datos")
    df_plot = df_long.copy()
    try: df_plot['Fecha_X'] = df_plot[col_periodo].dt.to_timestamp()
    except: df_plot['Fecha_X'] = pd.to_datetime(df_plot[col_periodo].astype(str), errors='coerce')
    df_plot = df_plot.dropna(subset=['Fecha_X']).sort_values('Fecha_X')
    df_plot['Mes_Texto'] = df_plot[col_periodo].apply(formatear_periodo_es)
    if es_porcentaje: df_plot['Etiqueta'] = df_plot[col_y].apply(lambda x: f"{x:.1f}%")
    else: df_plot['Etiqueta'] = df_plot[col_y].astype(str)
    color_map = {'Masculino': 'blue', 'Femenino': 'purple'} if (es_porcentaje and col_color) else None
    fig = px.line(df_plot, x='Fecha_X', y=col_y, color=col_color, title=titulo, markers=True, text='Etiqueta',
                  color_discrete_map=color_map, hover_data={'Fecha_X': False, 'Mes_Texto': True})
    fig.update_traces(textposition="top center")
    if es_porcentaje: fig.update_yaxes(range=[0, 115])
    return fig

def generar_grafica_plotly_circulos_edad(df_data, titulo):
    if df_data.empty: return px.scatter(title="Sin datos")
    df_data['Texto_Pct'] = df_data['Porcentaje'].apply(lambda x: f"{x:.0f}%")
    fig = px.scatter(df_data, x="Colonia", y="Rango", size="Porcentaje", text="Texto_Pct", title=titulo,
                     color_discrete_sequence=['black'], opacity=0.9)
    fig.update_traces(mode='markers+text', textposition='middle center', textfont=dict(color='white', weight='bold'))
    fig.update_yaxes(categoryorder='array', categoryarray=["Menor (0-17)", "Joven (18-29)", "Adulto (30-59)", "Mayor (60+)"])
    return fig

# --- ANÁLISIS ---

def analizar_lluvias_manual(df, col_lluvias, col_colonias, col_inc):
    if df.empty: return None
    df_l = df.copy()
    
    # CORRECCIÓN: Ignorar NaN en la columna de lluvias antes de procesar
    df_l = df_l.dropna(subset=[col_lluvias])
    
    df_l[col_lluvias] = df_l[col_lluvias].astype(str).str.lower().str.strip()
    positivos = ['sí', 'si', 'yes', 'true', '1', 'afirmativo', 'lluvia']
    df_l = df_l[df_l[col_lluvias].isin(positivos)]
    
    if df_l.empty: return None
    
    # CORRECCIÓN: Ignorar nulos en colonias o incidentes para el conteo específico
    conteo_col = df_l[col_colonias].dropna().value_counts()
    conteo_inc = df_l[col_inc].dropna().value_counts()
    
    return {
        'df_filtrado': df_l,
        'conteo_colonias': conteo_col,
        'conteo_incidentes': conteo_inc,
        'estadisticas': {'total_lluvias': len(df_l), 'porcentaje': (len(df_l)/len(df))*100}
    }

# (Se mantienen generar_reporte_word, generar_txt, get_link igual al original)
# ... [Omitido por brevedad] ...

def generar_reporte_word(conteos, imagenes):
    doc = Document()
    doc.add_heading('Reporte de Urgencias Operativas', 0).alignment = 1
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_heading('Resumen de Datos', 1)
    for nombre, conteo in conteos.items():
        if conteo is None or conteo.empty: continue
        doc.add_heading(nombre, 2)
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = 'Table Grid'
        tabla.rows[0].cells[0].text = "Categoría"
        tabla.rows[0].cells[1].text = "Cantidad"
        for k, v in list(conteo.items())[:20]:
            row = tabla.add_row().cells
            row[0].text = str(k).title()
            row[1].text = str(v)
    doc.add_page_break()
    doc.add_heading('Gráficas Visuales', 1)
    orden = ['General', 'Rango de Edad Dominante por Colonia', 'Padecimientos Detectados (Texto)',
             'Tendencia Porcentaje Género', 'Tendencia Incidentes', 'Tendencia Colonias', 
             'Tipos Lluvia', 'Colonias Lluvia', 'Tendencia Tipos Lluvia', 'Tendencia Total Lluvias',
             'Opiniones Técnicas por Clasificación', 'Tendencia Opiniones Técnicas por Mes']
    for titulo in orden:
        if titulo in imagenes and os.path.exists(imagenes[titulo]):
             doc.add_heading(titulo, 2)
             try: doc.add_picture(imagenes[titulo], width=Inches(6.5))
             except: doc.add_paragraph("[Error imagen]")
             doc.add_paragraph()
    out = os.path.join(tempfile.gettempdir(), 'reporte.docx')
    doc.save(out)
    return out

def generar_txt(conteos):
    txt = [f"REPORTE {datetime.now()}", "="*30]
    for k, v in conteos.items():
        if v is None: continue
        txt.append(f"\n{k.upper()}\n{'-'*len(k)}")
        for i, j in v.items(): txt.append(f"{i}: {j}")
    path = os.path.join(tempfile.gettempdir(), 'reporte.txt')
    with open(path, 'w', encoding='utf-8') as f: f.write("\n".join(txt))
    return path

def get_link(path, label):
    with open(path, "rb") as f: b64 = base64.b64encode(f.read()).decode()
    return f'<a href="data:file/octet;base64,{b64}" download="{os.path.basename(path)}">📥 {label}</a>'

# --- MAIN ---

def main():
    st.title("📊 Analizador de Urgencias Operativas")
    
    st.header("1. Datos Operativos (Principal)")
    f = st.file_uploader("Archivo Operativo (CSV/Excel)", type=['csv','xlsx'])
    if not f: 
        st.info("Sube un archivo para comenzar.")
        return

    try:
        df = pd.read_excel(f) if f.name.endswith('.xlsx') else pd.read_csv(f)
        st.success(f"Cargado: {len(df)} registros.")
    except Exception as e:
        st.error(f"Error leyendo archivo: {e}")
        return

    # 2. CONFIGURACIÓN (Se mantiene igual)
    st.header("2. Configuración General")
    c1, c2 = st.columns(2)
    col_inc = c1.selectbox("Columna INCIDENTES:", df.columns)
    col_col = c2.selectbox("Columna COLONIAS:", df.columns)
    
    c3, c4 = st.columns(2)
    col_genero = c3.selectbox("Columna GÉNERO (Opcional):", ["No usar"] + list(df.columns))
    if col_genero == "No usar": col_genero = None
    
    col_edad = c4.selectbox("Columna EDAD (Opcional):", ["No usar"] + list(df.columns))
    if col_edad == "No usar": col_edad = None

    col_fecha = st.selectbox("Columna FECHAS (Necesaria para gráficas de línea):", ["No usar"]+list(df.columns))
    if col_fecha == "No usar": col_fecha = None
    
    st.subheader("🌧️ Análisis de Lluvias")
    check_lluvias = st.checkbox("Analizar Reportes de Lluvias")
    col_lluvias = None
    if check_lluvias:
        col_lluvias = st.selectbox("Columna Indicador Lluvia (Sí/No, 1/0):", df.columns)

    st.markdown("---")
    st.subheader("🏥 Análisis de Descripción (Padecimientos/Enfermedades)")
    check_txt_med = st.checkbox("Analizar Descripciones (Ej. Buscar Diabetes, Hipertensión, etc.)")
    
    col_desc_med = None
    col_filtro_med = None
    val_filtro_med = None
    
    if check_txt_med:
        c_txt1, c_txt2 = st.columns(2)
        idx_desc = 0
        if "descripcion_del_incidente" in df.columns:
            idx_desc = list(df.columns).index("descripcion_del_incidente")
            
        col_desc_med = c_txt1.selectbox("Columna con la DESCRIPCIÓN:", df.columns, index=idx_desc)
        col_filtro_med = c_txt2.selectbox("Columna para FILTRAR (Ej. Tipo de Incidente):", df.columns, index=0)
        
        unique_vals = df[col_filtro_med].dropna().unique().tolist()
        val_filtro_med = st.multiselect(f"Selecciona valores de '{col_filtro_med}' a analizar (Ej. Atencion Medica):", unique_vals)
        
        if not val_filtro_med:
            st.warning("⚠️ Debes seleccionar al menos un valor para filtrar.")

    # --- NUEVA SECCIÓN: OPINIONES TÉCNICAS ---
    st.markdown("---")
    st.subheader("📂 Módulo Extra: Opiniones Técnicas")
    check_opiniones = st.checkbox("Analizar Opiniones Técnicas (Subir otro documento)")
    
    df_op = None
    col_fecha_op = None
    col_tipo_op = None
    col_colonia_op = None
    
    if check_opiniones:
        f_op = st.file_uploader("Subir Archivo de Opiniones Técnicas (CSV/Excel)", type=['csv','xlsx'])
        if f_op:
            try:
                df_op = pd.read_excel(f_op) if f_op.name.endswith('.xlsx') else pd.read_csv(f_op)
                st.success(f"Cargado Opiniones: {len(df_op)} registros.")
                
                c_op1, c_op2, c_op3 = st.columns(3)
                col_fecha_op = c_op1.selectbox("Columna FECHA (Opiniones):", df_op.columns)
                col_tipo_op = c_op2.selectbox("Columna CLASIFICACIÓN (Ej. Arbolado):", df_op.columns)
                col_colonia_op = c_op3.selectbox("Columna COLONIA (Opiniones):", df_op.columns)
                
            except Exception as e:
                st.error(f"Error archivo opiniones: {e}")

    st.markdown("---")
    st.subheader("🛠️ Filtros y Gráficas Avanzadas")
    ignorar_medica = st.checkbox("Ignorar 'Atención Médica' (En gráfica General)", value=True)
    
    col_g1, col_g2 = st.columns(2)
    graf_top10 = col_g1.checkbox("Barras: Top 10 Incidentes", value=True)
    graf_pct_genero = col_g2.checkbox("Línea: Porcentaje Género", value=False)
    
    graf_linea_inc = col_g1.checkbox("Línea: Tendencia Top 10 Incidentes", value=False)
    graf_linea_col = col_g2.checkbox("Línea: Tendencia Top 10 Colonias", value=False)
    
    graf_edad_colonia = col_g1.checkbox("Círculos: Rango de Edad Dominante por Colonia", value=False)

    if (graf_linea_inc or graf_linea_col or graf_pct_genero) and not col_fecha:
         st.warning("⚠️ Las gráficas de línea requieren una Columna de FECHAS.")
    if graf_pct_genero and not col_genero:
         st.warning("⚠️ Requiere Columna GÉNERO.")
    if graf_edad_colonia and not col_edad:
         st.warning("⚠️ Requiere Columna EDAD.")

    # 3. PROCESAR
    if st.button("🚀 Generar Reporte", type="primary"):
        with st.spinner("Procesando..."):
            df_c = df.copy()
            df_c[col_inc] = df_c[col_inc].apply(limpiar_texto)
            df_c[col_col] = df_c[col_col].apply(limpiar_texto)
            
            df_general = df_c.copy()
            if ignorar_medica:
                df_general = df_general[df_general[col_inc] != "atencion medica"]
            
            if df_general.empty:
                st.warning("Sin datos para gráfica general tras filtros (pero seguimos procesando lo demás).")

            valid_fechas = False
            if col_fecha:
                df_c['fecha_p'] = df_c[col_fecha].apply(parsear_fecha)
                df_general['fecha_p'] = df_c['fecha_p']
                
                if df_c['fecha_p'].notna().sum() > 0:
                    valid_fechas = True
                    df_c = df_c.dropna(subset=['fecha_p'])
                    df_general = df_general.dropna(subset=['fecha_p'])
                    df_c['mes'] = df_c['fecha_p'].dt.to_period('M')
                    df_general['mes'] = df_general['fecha_p'].dt.to_period('M')
            
            conteos = {
                "General": df_general[col_inc].value_counts(),
                "Colonias": df_general[col_col].value_counts()
            }
            imgs = {}

            # PROCESAR OPINIONES TÉCNICAS
            if check_opiniones and df_op is not None and col_fecha_op and col_tipo_op:
                try:
                    df_op_c = df_op.copy()
                    # Limpiamos nulos antes de procesar texto
                    df_op_c = df_op_c.dropna(subset=[col_tipo_op])
                    df_op_c[col_tipo_op] = df_op_c[col_tipo_op].apply(limpiar_texto)
                    
                    # 1. Gráfica de Barras (Acumulado)
                    conteo_tipos_op = df_op_c[col_tipo_op].value_counts()
                    conteos["Opiniones Técnicas por Clasificación"] = conteo_tipos_op
                    
                    st.subheader("📋 Opiniones Técnicas: Clasificación")
                    st.plotly_chart(generar_grafica_plotly_bar(conteo_tipos_op, "Clasificación Opiniones"), use_container_width=True)
                    imgs["Opiniones Técnicas por Clasificación"] = generar_grafica_bar(conteo_tipos_op, "Opiniones Técnicas por Clasificación", "g_op_bar.png")
                    
                    # 2. Gráfica de Línea (Tendencia)
                    df_op_c['fecha_p'] = df_op_c[col_fecha_op].apply(parsear_fecha)
                    df_op_c['fecha_p'] = pd.to_datetime(df_op_c['fecha_p'], errors='coerce')
                    df_op_c = df_op_c.dropna(subset=['fecha_p'])

                    if not df_op_c.empty:
                        df_op_c['mes'] = df_op_c['fecha_p'].dt.to_period('M')
                        data_linea_op = df_op_c.groupby(['mes', col_tipo_op]).size().reset_index(name='conteo')
                        
                        st.subheader("📈 Opiniones Técnicas: Tendencia Mensual")
                        st.plotly_chart(generar_grafica_plotly_linea(data_linea_op, 'mes', 'conteo', col_tipo_op, "Evolución Opiniones"), use_container_width=True)
                        imgs["Tendencia Opiniones Técnicas por Mes"] = generar_grafica_linea_multiple(data_linea_op, 'mes', 'conteo', col_tipo_op, "Evolución Opiniones Técnicas", "l_op_time.png")
                    else:
                        st.warning("No se pudieron detectar fechas válidas en el archivo de Opiniones.")
                        
                except Exception as e:
                    st.error(f"Error procesando opiniones: {e}")

            if check_txt_med and col_desc_med and col_filtro_med and val_filtro_med:
                try:
                    df_med = df[df[col_filtro_med].isin(val_filtro_med)].copy()
                    if not df_med.empty:
                        df_med['Padecimiento_Detectado'] = df_med[col_desc_med].apply(clasificar_enfermedad)
                        
                        conteo_padecimientos = df_med['Padecimiento_Detectado'].str.split(', ', expand=True).stack().value_counts()
                        
                        conteos["Padecimientos Detectados (Texto)"] = conteo_padecimientos
                        
                        st.subheader(f"🏥 Padecimientos Detectados en {val_filtro_med}")
                        st.plotly_chart(generar_grafica_plotly_bar(conteo_padecimientos.head(15), "Top Padecimientos Detectados"), use_container_width=True)
                        imgs["Padecimientos Detectados (Texto)"] = generar_grafica_bar(conteo_padecimientos.head(15), "Padecimientos Detectados", "g_pad_txt.png")
                    else:
                        st.warning("No se encontraron filas con los filtros seleccionados para el análisis de texto.")
                except Exception as e:
                    st.error(f"Error en análisis de texto: {e}")

            res_lluv = None
            if check_lluvias and col_lluvias:
                res_lluv = analizar_lluvias_manual(df_c, col_lluvias, col_col, col_inc)
                if res_lluv:
                    conteos["Tipos de Incidentes en Lluvias"] = res_lluv['conteo_incidentes']
                    conteos["Colonias Afectadas por Lluvias"] = res_lluv['conteo_colonias']
            
            if col_genero:
                 df_general['genero_norm'] = df_general[col_genero].apply(normalizar_genero)
                 conteos["Desglose por Género"] = df_general['genero_norm'].fillna('No id').value_counts()

            if col_edad:
                df_general['edad_cat'] = df_general[col_edad].apply(limpiar_y_categorizar_edad)
                conteos["Desglose por Rango Edad"] = df_general['edad_cat'].value_counts()

            st.header("3. Resultados")
            c1, c2, c3 = st.columns(3)
            c1.metric("Total", len(df_c))
            c2.metric("Tipos", df_general[col_inc].nunique())
            if res_lluv: c3.metric("Lluvia", res_lluv['estadisticas']['total_lluvias'])
            
            st.subheader("Vista General")
            top_gen = conteos["General"].head(15)
            st.plotly_chart(generar_grafica_plotly_bar(top_gen, "Top Incidentes"), use_container_width=True)
            imgs["General"] = generar_grafica_bar(top_gen, "Top Incidentes", "g1.png")

            # ... Resto de la lógica de gráficas se mantiene igual ...
            if col_edad and graf_edad_colonia:
                st.subheader("⚫ Rango de Edad Dominante por Colonia (Top 10)")
                try:
                    top10_c = conteos["Colonias"].head(10).index.tolist()
                    df_edad = df_general[df_general[col_col].isin(top10_c) & df_general['edad_cat'].notna()].copy()
                    
                    if not df_edad.empty:
                        grp_edad = df_edad.groupby([col_col, 'edad_cat']).size().reset_index(name='Cantidad')
                        total_por_colonia = df_edad.groupby(col_col).size().reset_index(name='Total_Col')
                        grp_edad = pd.merge(grp_edad, total_por_colonia, on=col_col)
                        grp_edad['Porcentaje'] = (grp_edad['Cantidad'] / grp_edad['Total_Col']) * 100
                        grp_edad.rename(columns={col_col: 'Colonia', 'edad_cat': 'Rango'}, inplace=True)
                        
                        grp_edad = grp_edad.sort_values(['Colonia', 'Porcentaje'], ascending=[True, False])
                        grp_edad_max = grp_edad.drop_duplicates(subset=['Colonia'], keep='first')
                        
                        st.plotly_chart(generar_grafica_plotly_circulos_edad(grp_edad_max, "Rango de Edad Dominante por Colonia"), use_container_width=True)
                        imgs["Rango de Edad Dominante por Colonia"] = generar_grafica_circulos_edad_word(grp_edad_max, "Rango de Edad Dominante (Círculo Negro = % Mayor)", "g_edad_circ.png")
                    else:
                        st.warning("No hay datos de edad suficientes en las top 10 colonias.")
                except Exception as e:
                    st.error(f"Error en gráfica de edades: {e}")

            if valid_fechas and col_genero and graf_pct_genero:
                try:
                    df_gen = df_general[df_general['genero_norm'].isin(['Masculino', 'Femenino'])].copy()
                    if not df_gen.empty:
                        conteo_gen_mes = df_gen.groupby(['mes', 'genero_norm']).size().reset_index(name='cuenta')
                        total_mes = df_gen.groupby('mes').size().reset_index(name='total_mes')
                        df_pct_gen = pd.merge(conteo_gen_mes, total_mes, on='mes')
                        df_pct_gen['porcentaje'] = (df_pct_gen['cuenta'] / df_pct_gen['total_mes']) * 100
                        
                        st.subheader("📈 Género (Hombres vs Mujeres)")
                        st.plotly_chart(generar_grafica_plotly_linea(df_pct_gen, 'mes', 'porcentaje', 'genero_norm', "Evolución % Género", True), use_container_width=True)
                        imgs["Tendencia Porcentaje Género"] = generar_grafica_linea_porcentaje_genero(df_pct_gen, 'mes', 'porcentaje', 'genero_norm', "Solicitantes por Género (%)", "l_pct_gen.png")
                    else:
                         st.warning("Datos de género insuficientes.")
                except Exception as e: st.error(f"Error género: {e}")
            
            if valid_fechas:
                if graf_linea_inc:
                    try:
                        top10_names = conteos["General"].head(10).index.tolist()
                        df_top = df_general[df_general[col_inc].isin(top10_names)].copy()
                        data_linea = df_top.groupby(['mes', col_inc]).size().reset_index(name='conteo')
                        if not data_linea.empty:
                            st.subheader("📈 Tendencia Incidentes")
                            st.plotly_chart(generar_grafica_plotly_linea(data_linea, 'mes', 'conteo', col_inc, "Evolución Incidentes"), use_container_width=True)
                            imgs["Tendencia Incidentes"] = generar_grafica_linea_multiple(data_linea, 'mes', 'conteo', col_inc, "Comparativa Incidentes", "l_inc.png")
                    except: pass
                
                if graf_linea_col:
                    try:
                        top10_cols = conteos["Colonias"].head(10).index.tolist()
                        df_top_c = df_general[df_general[col_col].isin(top10_cols)].copy()
                        data_linea_c = df_top_c.groupby(['mes', col_col]).size().reset_index(name='conteo')
                        if not data_linea_c.empty:
                            st.subheader("📈 Tendencia Colonias")
                            st.plotly_chart(generar_grafica_plotly_linea(data_linea_c, 'mes', 'conteo', col_col, "Evolución Colonias"), use_container_width=True)
                            imgs["Tendencia Colonias"] = generar_grafica_linea_multiple(data_linea_c, 'mes', 'conteo', col_col, "Comparativa Colonias", "l_col.png")
                    except: pass

            if res_lluv:
                st.markdown("---")
                st.header("🌧️ Análisis Lluvias")
                top_inc_lluv = res_lluv['conteo_incidentes'].head(15)
                st.plotly_chart(generar_grafica_plotly_bar(top_inc_lluv, "Tipos (Lluvias)"), use_container_width=True)
                imgs["Tipos Lluvia"] = generar_grafica_bar(top_inc_lluv, "Tipos en Lluvias", "g_inc_lluv.png")
                
                top_col_lluv = res_lluv['conteo_colonias'].head(15)
                st.plotly_chart(generar_grafica_plotly_bar(top_col_lluv, "Colonias (Lluvias)"), use_container_width=True)
                imgs["Colonias Lluvia"] = generar_grafica_bar(top_col_lluv, "Colonias en Lluvias", "g_col_lluv.png")
                
                if valid_fechas:
                    df_lluvia_t = res_lluv['df_filtrado'].copy()
                    df_lluvia_t['fecha_p'] = pd.to_datetime(df_lluvia_t['fecha_p'], errors='coerce')
                    df_lluvia_t = df_lluvia_t.dropna(subset=['fecha_p'])
                    
                    if not df_lluvia_t.empty:
                        df_lluvia_t['mes'] = df_lluvia_t['fecha_p'].dt.to_period('M')
                        try:
                            top5_inc_lluvia = res_lluv['conteo_incidentes'].head(5).index.tolist()
                            df_top_lluvia = df_lluvia_t[df_lluvia_t[col_inc].isin(top5_inc_lluvia)]
                            if not df_top_lluvia.empty:
                                data_linea_lluvia = df_top_lluvia.groupby(['mes', col_inc]).size().reset_index(name='conteo')
                                st.subheader("📈 Tendencia Tipos (Lluvias)")
                                st.plotly_chart(generar_grafica_plotly_linea(data_linea_lluvia, 'mes', 'conteo', col_inc, "Evolución Tipos (Lluvias)"), use_container_width=True)
                                imgs["Tendencia Tipos Lluvia"] = generar_grafica_linea_multiple(data_linea_lluvia, 'mes', 'conteo', col_inc, "Evolución Tipos (Lluvias)", "l_tipo_lluv.png")
                        except: pass
                        
                        try:
                            data_total = df_lluvia_t.groupby('mes').size().reset_index(name='conteo')
                            if not data_total.empty:
                                imgs["Tendencia Total Lluvias"] = generar_grafica_linea_simple(data_total.set_index('mes')['conteo'], "Total Lluvias", "Mes", "Cant", "l_total_lluv.png")
                        except: pass

            st.success("✅ Hecho")
            c1, c2 = st.columns(2)
            c1.markdown(get_link(generar_reporte_word(conteos, imgs), "Word"), unsafe_allow_html=True)
            c2.markdown(get_link(generar_txt(conteos), "Txt"), unsafe_allow_html=True)
            for p in imgs.values(): 
                if p and os.path.exists(p): os.remove(p)

if __name__ == "__main__":
    main()
